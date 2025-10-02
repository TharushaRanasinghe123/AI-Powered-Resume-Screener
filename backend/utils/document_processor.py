import pdfplumber
from docx import Document
import os
import re
from typing import Dict, Any
import magic

class DocumentProcessor:
    """
    Handles text extraction from various file formats.
    """

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> Dict[str, Any]:
        """Extract text from a PDF files."""
        try:
            full_text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        full_text += text + "\n"
            
            return{
                "success": True,
                "text": full_text.strip(),
                "page_count": len(pdf.pages),
                "file_type": "pdf"
            }
        
        except Exception as e:
            return{
                "success": False,
                "error": f"PDF extraction failed: {str(e)}",
                "text": "",
                "file_type": "pdf"
            }
        
    
    @staticmethod
    def extract_from_docx(file_path: str) -> Dict[str, Any]:
        """ Extract text from DOCX files."""

        try: 
            doc = Document(file_path)
            full_text = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)

            return{
                "success": True,
                "text": "\n".join(full_text),
                "file_type": "docx"
            }
        except Exception as e:
            return{
                "success": False,
                "error": f"DOCX extraction failed: {str(e)}",
                "text": "",
                "file_type": "docx"
            }


                